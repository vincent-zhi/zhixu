"""Parse user-uploaded docs (md/txt/pdf/docx) into a normalized JSON payload.

Usage:
    python parse_user_docs.py --files path1 [path2 ...]

Output (stdout):
    {
      "documents": [
        {"path": "...", "type": "md", "text": "...", "tables": [...], "inherited_images": [...]},
        ...
      ],
      "errors": [...]
    }

Tables: each item = {"doc_index": int, "table_index": int, "rows": [[str, ...]]}
Inherited images: each item = {"doc_index": int, "image_index": int, "path": str, "alt": str}
                  (path is absolute; images embedded in docx/pdf are extracted
                  to the same directory as the source doc under
                  `<docname>_inherited/`; md image references are resolved relative
                  to the md file's directory).

No extra deps beyond stdlib + pypdf + python-docx. Markdown-first (table + image
extraction), pdf/docx best-effort.
"""
from __future__ import annotations

import argparse
import json
import re
import sys
from pathlib import Path

from docx.opc.exceptions import PackageNotFoundError
from pypdf.errors import PdfReadError

MAX_CHARS = 20000


# ---------------------------------------------------------------------------
# Text truncation
# ---------------------------------------------------------------------------


def _truncate(text: str) -> str:
    if len(text) <= MAX_CHARS:
        return text
    return text[:MAX_CHARS] + "\n[TRUNCATED]"


# ---------------------------------------------------------------------------
# Markdown parser: tables + image references
# ---------------------------------------------------------------------------


_MD_IMG_RE = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
_MD_TABLE_SEPARATOR_RE = re.compile(r"^\s*\|?\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$")


def _split_md_table_row(line: str) -> list[str]:
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [cell.strip() for cell in s.split("|")]


def _extract_md_tables(text: str) -> list[list[list[str]]]:
    """Find GitHub-flavored markdown tables. Returns list of tables;
    each table is a list of rows; each row is a list of cells."""
    lines = text.splitlines()
    tables: list[list[list[str]]] = []
    i = 0
    while i < len(lines):
        # A table needs: header row (with |), separator row (|---|---|), data rows.
        if "|" in lines[i] and i + 1 < len(lines) and _MD_TABLE_SEPARATOR_RE.match(lines[i + 1]):
            header = _split_md_table_row(lines[i])
            rows: list[list[str]] = [header]
            j = i + 2
            while j < len(lines) and "|" in lines[j] and lines[j].strip():
                rows.append(_split_md_table_row(lines[j]))
                j += 1
            if len(rows) >= 2:
                tables.append(rows)
            i = j
        else:
            i += 1
    return tables


def _extract_md_images(text: str, base_dir: Path) -> list[dict]:
    """Find `![alt](path)` references. Resolve relative paths against base_dir.
    Leave http(s) URLs as-is. Return absolute path for local files if they
    exist on disk; else skip (we can't download remote, but keep the URL)."""
    out: list[dict] = []
    for m in _MD_IMG_RE.finditer(text):
        alt = m.group(1).strip()
        target = m.group(2).strip()
        if target.startswith(("http://", "https://", "data:")):
            out.append({"path": target, "alt": alt})
            continue
        # strip any ?query or #fragment or spaces
        target_clean = target.split()[0].split("?")[0].split("#")[0]
        p = Path(target_clean)
        if not p.is_absolute():
            p = (base_dir / p).resolve()
        if p.exists() and p.is_file():
            out.append({"path": str(p), "alt": alt})
    return out


def parse_md_or_txt(path: Path) -> dict:
    raw = path.read_text(encoding="utf-8", errors="replace")
    ext = path.suffix.lstrip(".").lower()
    tables = _extract_md_tables(raw) if ext == "md" else []
    images = _extract_md_images(raw, path.parent) if ext == "md" else []
    return {
        "path": str(path),
        "type": ext,
        "text": _truncate(raw),
        "tables": tables,
        "inherited_images": images,
    }


# ---------------------------------------------------------------------------
# PDF parser: text + best-effort image extraction
# ---------------------------------------------------------------------------


def parse_pdf(path: Path) -> dict:
    from pypdf import PdfReader
    try:
        reader = PdfReader(str(path))
        pages = [p.extract_text() or "" for p in reader.pages]
    except PdfReadError as exc:
        raise ValueError(f"corrupted pdf: {exc}") from exc
    text = "\n".join(pages)

    # Best-effort image extraction. pypdf exposes page.images since 3.x;
    # if it fails or there are none, we silently skip.
    images: list[dict] = []
    try:
        out_dir = path.parent / f"{path.stem}_inherited"
        idx = 0
        for page_no, page in enumerate(reader.pages, start=1):
            try:
                page_images = list(page.images)
            except Exception:
                continue
            for img in page_images:
                try:
                    name = getattr(img, "name", None) or f"p{page_no}_i{idx}"
                    ext = Path(name).suffix or ".png"
                    out_dir.mkdir(parents=True, exist_ok=True)
                    out_path = out_dir / f"page{page_no:03d}_{idx:02d}{ext}"
                    with out_path.open("wb") as f:
                        f.write(img.data)
                    images.append({"path": str(out_path), "alt": f"pdf page {page_no}"})
                    idx += 1
                except Exception:
                    continue
    except Exception:
        pass  # image extraction is best-effort

    # Note: pypdf doesn't give structured tables — tables stay embedded in text.
    return {
        "path": str(path),
        "type": "pdf",
        "text": _truncate(text),
        "pages": len(reader.pages),
        "tables": [],  # best-effort pdf table extraction out of scope (no pdfplumber)
        "inherited_images": images,
    }


# ---------------------------------------------------------------------------
# DOCX parser: text + tables (native) + embedded images
# ---------------------------------------------------------------------------


def parse_docx(path: Path) -> dict:
    from docx import Document
    try:
        doc = Document(str(path))
    except PackageNotFoundError as exc:
        raise ValueError(f"corrupted or invalid docx: {exc}") from exc

    # Text = paragraphs (keep order). Tables are separate channel now.
    paragraphs = [p.text for p in doc.paragraphs if p.text]
    text = "\n".join(paragraphs)

    # Tables: native structured extraction
    tables: list[list[list[str]]] = []
    for t in doc.tables:
        rows = [[cell.text for cell in row.cells] for row in t.rows]
        if rows:
            tables.append(rows)

    # Embedded images: iterate part.related_parts and save each `image/*` blob
    images: list[dict] = []
    try:
        out_dir = path.parent / f"{path.stem}_inherited"
        for idx, (_, rel) in enumerate(doc.part.related_parts.items()):
            content_type = getattr(rel, "content_type", "") or ""
            if not content_type.startswith("image/"):
                continue
            try:
                blob = rel.blob
                # guess extension from content_type
                ext = "." + content_type.split("/", 1)[1].split(";")[0].split("+")[0]
                if ext == ".jpeg":
                    ext = ".jpg"
                if ext not in {".png", ".jpg", ".gif", ".webp", ".bmp"}:
                    ext = ".png"
                out_dir.mkdir(parents=True, exist_ok=True)
                out_path = out_dir / f"image_{idx:02d}{ext}"
                with out_path.open("wb") as f:
                    f.write(blob)
                images.append({"path": str(out_path), "alt": f"docx image {idx}"})
            except Exception:
                continue
    except Exception:
        pass

    return {
        "path": str(path),
        "type": "docx",
        "text": _truncate(text),
        "paragraphs": len(paragraphs),
        "tables": tables,
        "inherited_images": images,
    }


# ---------------------------------------------------------------------------
# Dispatch + main
# ---------------------------------------------------------------------------


def dispatch(path: Path) -> dict:
    ext = path.suffix.lower()
    if ext in (".md", ".txt"):
        return parse_md_or_txt(path)
    if ext == ".pdf":
        return parse_pdf(path)
    if ext == ".docx":
        return parse_docx(path)
    raise ValueError(f"unsupported type: {ext}")


def main(argv: list[str] | None = None) -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--files", nargs="+", required=True)
    parser.add_argument(
        "--output",
        type=str,
        default=None,
        help=(
            "Write the JSON result to this file instead of stdout. "
            "Parent directories are created if missing. Recommended when the "
            "caller is an agent that may not reliably handle shell redirection."
        ),
    )
    args = parser.parse_args(argv)
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")

    documents: list[dict] = []
    errors: list[dict] = []
    for raw in args.files:
        p = Path(raw).expanduser().resolve()
        if not p.exists():
            errors.append({"path": str(p), "error": "file not found"})
            continue
        if not p.is_file():
            errors.append({"path": str(p), "error": "not a regular file"})
            continue
        try:
            documents.append(dispatch(p))
        except Exception as exc:  # noqa: BLE001
            errors.append({"path": str(p), "error": f"{type(exc).__name__}: {exc}"})

    # Annotate indices for cross-reference from digest / outline
    for di, d in enumerate(documents):
        d["doc_index"] = di
        for ti, _t in enumerate(d.get("tables") or []):
            pass  # kept as plain list for simplicity; consumers index by position
        for ii, img in enumerate(d.get("inherited_images") or []):
            img["image_index"] = ii
            img["doc_index"] = di

    payload = {"documents": documents, "errors": errors}
    if args.output:
        out_path = Path(args.output).expanduser().resolve()
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(
            json.dumps(payload, ensure_ascii=False),
            encoding="utf-8",
        )
        # Emit a short status line so the agent has something to echo.
        print(json.dumps({
            "status": "ok",
            "output": str(out_path),
            "documents": len(documents),
            "errors": len(errors),
        }, ensure_ascii=False))
    else:
        json.dump(payload, sys.stdout, ensure_ascii=False)
    return 0


if __name__ == "__main__":
    sys.exit(main())
